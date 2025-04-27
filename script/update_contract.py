# 

import os
import json
from docx import Document
from pathlib import Path

from docx import Document

def extract_text_from_docx(file_path):
    """提取 docx 富文本为结构化 HTML 字符串"""
    try:
        doc = Document(file_path)
        body_content = []

        def format_run(run):
            text = run.text
            if not text.strip():
                return ""
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"

            styles = []
            if run.font.name:
                styles.append(f"font-family:'{run.font.name}'")
            if run.font.size:
                styles.append(f"font-size:{run.font.size.pt}pt")
            if run.font.color and run.font.color.rgb:
                styles.append(f"color:#{run.font.color.rgb}")
            style_str = "; ".join(styles)

            return f"<span style=\"{style_str}\">{text}</span>" if style_str else text

        for para in doc.paragraphs:
            if para.text.strip():
                formatted = ''.join([format_run(run) for run in para.runs])
                align = para.alignment.name if para.alignment else "left"
                body_content.append(f"<p style=\"text-align:{align.lower()};\">{formatted}</p>")

        for table in doc.tables:
            table_html = ["<table style=\"border-collapse:collapse; width:100%;\" border=\"1\">"]
            for row in table.rows:
                table_html.append("<tr>")
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        para_html = ''.join([format_run(run) for run in para.runs])
                        cell_text.append(f"<p>{para_html}</p>")
                    table_html.append(f"<td>{''.join(cell_text)}</td>")
                table_html.append("</tr>")
            table_html.append("</table>")
            body_content.append(''.join(table_html))

        # 加上标准 HTML 头
        html = f"""<html>
<head>
<meta http-equiv="Content-Type" content="text/html; charset=utf-8" />
<meta http-equiv="Content-Style-Type" content="text/css" />
<title>合同文档</title>
</head>
<body>
<div>
{''.join(body_content)}
</div>
</body>
</html>"""
        return html

    except Exception as e:
        print(f"❌ 处理文件 {file_path} 时出错: {e}")
        return ""


def update_json_with_fulltext():
    """更新vector_new.json文件，为每条数据添加fulltext字段"""
    template_dir = "/home/user/opt/ssy/contract_template/data/修改过的文件"
    json_file_path = "/home/user/opt/ssy/contract_template/data/vector_data/vector_new.json"
    
    docx_files = {}
    if os.path.exists(template_dir):
        for file in os.listdir(template_dir):
            if file.endswith('.docx'):
                template_name = file[:-5]
                docx_files[template_name] = os.path.join(template_dir, file)
        print(f"在模板目录中找到 {len(docx_files)} 个docx文件")
    else:
        print(f"模板目录 {template_dir} 不存在")
        return
    
    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"读取JSON文件时出错: {e}")
        return

    template_to_items = {}
    for item in data:
        template_name = item.get('template', '')
        if template_name:
            if template_name not in template_to_items:
                template_to_items[template_name] = []
            template_to_items[template_name].append(item)
    
    updated_count = 0
    for template_name, file_path in docx_files.items():
        if template_name in template_to_items:
            fulltext = extract_text_from_docx(file_path)
            for item in template_to_items[template_name]:
                item['fulltext'] = fulltext
                updated_count += 1
            print(f"已更新模板 '{template_name}' 的fulltext字段，共 {len(template_to_items[template_name])} 条数据")
        else:
            print(f"模板 '{template_name}' 在JSON数据中没有匹配项")
    
    # for template_name in template_to_items:
    #     if template_name not in docx_files:
    #         for item in template_to_items[template_name]:
    #             item['fulltext'] = ""
    #         print(f"模板 '{template_name}' 在模板目录中没有对应的docx文件，已设置fulltext为空字符串")
    
    try:
        with open(json_file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"已成功更新 {json_file_path}，共更新 {updated_count} 条数据")
    except Exception as e:
        print(f"保存JSON文件时出错: {e}")

if __name__ == "__main__":
    update_json_with_fulltext()